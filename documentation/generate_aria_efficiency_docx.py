#!/usr/bin/env python3
"""
Generates a properly-formatted Word document for the ARIA Efficiency Whitepaper.
Includes: cover page, auto-TOC field, all sections with tables, charts, appendices.
Run:    python generate_aria_efficiency_docx.py
Output: documentation/ARIA_Efficiency_Whitepaper.docx
Requires: pip install python-docx matplotlib
"""

import io
import os
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import numpy as np

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Colour palette ─────────────────────────────────────────────────────────────
DARK_NAVY   = (0x00, 0x2B, 0x5C)
MID_BLUE    = (0x00, 0x5B, 0x99)
BRAND_BLUE  = (0x1D, 0x6F, 0xDB)
ACCENT_TEAL = (0x00, 0x7A, 0x87)
LIGHT_GREY  = (0xF2, 0xF4, 0xF7)
WHITE       = (0xFF, 0xFF, 0xFF)
DARK_TEXT   = (0x1A, 0x1A, 0x2E)
AMBER       = (0xE8, 0x8C, 0x00)
EMERALD     = (0x05, 0x96, 0x69)
ROSE        = (0xE5, 0x3E, 0x3E)

def rgb(t): return RGBColor(t[0], t[1], t[2])
def hex_color(t): return f"{t[0]:02X}{t[1]:02X}{t[2]:02X}"


# ── Low-level helpers ──────────────────────────────────────────────────────────

def set_cell_bg(cell, color_tuple):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color(color_tuple))
    tcPr.append(shd)

def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), val.get('val', 'single'))
            el.set(qn('w:sz'), str(val.get('sz', 4)))
            el.set(qn('w:color'), val.get('color', '000000'))
            tcBorders.append(el)
    tcPr.append(tcBorders)

def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(docx_break_type('page'))

def docx_break_type(btype):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    br = OxmlElement('w:br')
    br.set(qn('w:type'), btype)
    return br

def add_hr(doc, color=BRAND_BLUE, thickness=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(thickness))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), hex_color(color))
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def para_shade(p, color_tuple):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color(color_tuple))
    pPr.append(shd)

def insert_toc(doc):
    """Insert a Word TOC field that auto-populates on Ctrl+A, F9."""
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run()
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
    fldChar_separate = OxmlElement('w:fldChar')
    fldChar_separate.set(qn('w:fldCharType'), 'separate')
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar_begin)
    run._r.append(instrText)
    run._r.append(fldChar_separate)
    run._r.append(fldChar_end)


# ── Typography helpers ─────────────────────────────────────────────────────────

def style_run(run, size=11, bold=False, italic=False, color=DARK_TEXT, font='Calibri'):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = rgb(color)

def add_heading(doc, text, level=1, color=DARK_NAVY, size=None, before=12, after=6):
    sizes = {1: 18, 2: 14, 3: 12, 4: 11}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.style = f'Heading {level}'
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(size or sizes.get(level, 11))
    run.font.bold = True
    run.font.color.rgb = rgb(color)
    return p

def add_body(doc, text, color=DARK_TEXT, size=10.5, before=2, after=4, indent=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    style_run(run, size=size, color=color)
    return p

def add_bullet(doc, text, color=DARK_TEXT, size=10.5, indent_cm=0.5):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(indent_cm)
    run = p.add_run(text)
    style_run(run, size=size, color=color)
    return p

def add_callout(doc, text, bg=LIGHT_GREY, color=DARK_NAVY, size=11, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    para_shade(p, bg)
    run = p.add_run(text)
    style_run(run, size=size, color=color, bold=bold)
    return p


# ── Table helpers ──────────────────────────────────────────────────────────────

def add_table(doc, headers, rows, col_widths=None, header_bg=DARK_NAVY, row_alt_bg=LIGHT_GREY):
    n_cols = len(headers)
    tbl = doc.add_table(rows=1 + len(rows), cols=n_cols)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, header_bg)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(h)
        style_run(run, size=9, bold=True, color=WHITE)

    # Data rows
    for ri, row_data in enumerate(rows):
        row = tbl.rows[ri + 1]
        bg = row_alt_bg if ri % 2 == 1 else WHITE
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            set_cell_bg(cell, bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            style_run(run, size=9, color=DARK_TEXT)

    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()
    return tbl


# ── Chart helpers ──────────────────────────────────────────────────────────────

def make_bar_chart(labels, values, title, ylabel='', color='#1D6FDB', figsize=(7, 3)):
    fig, ax = plt.subplots(figsize=figsize)
    ax.set_facecolor('#F2F4F7')
    fig.patch.set_facecolor('#FFFFFF')
    bars = ax.bar(labels, values, color=color, width=0.55, zorder=3)
    ax.set_title(title, fontsize=11, fontweight='bold', color='#002B5C', pad=10)
    ax.set_ylabel(ylabel, fontsize=9, color='#333333')
    ax.tick_params(axis='x', labelsize=8, rotation=15)
    ax.tick_params(axis='y', labelsize=8)
    ax.yaxis.grid(True, color='white', linewidth=1.2, zorder=0)
    ax.set_axisbelow(True)
    for spine in ax.spines.values():
        spine.set_visible(False)
    for bar, val in zip(bars, values):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + max(values)*0.01,
                str(val), ha='center', va='bottom', fontsize=8, color='#002B5C', fontweight='bold')
    fig.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150)
    plt.close(fig)
    buf.seek(0)
    return buf

def make_comparison_bar(categories, vals_ai, vals_team, title, figsize=(7, 3.5)):
    x = np.arange(len(categories))
    width = 0.35
    fig, ax = plt.subplots(figsize=figsize)
    ax.set_facecolor('#F2F4F7')
    fig.patch.set_facecolor('#FFFFFF')
    b1 = ax.bar(x - width/2, vals_ai, width, label='1 PO + Claude', color='#1D6FDB', zorder=3)
    b2 = ax.bar(x + width/2, vals_team, width, label='3-Person Team', color='#AAAACC', zorder=3)
    ax.set_title(title, fontsize=11, fontweight='bold', color='#002B5C', pad=10)
    ax.set_xticks(x)
    ax.set_xticklabels(categories, fontsize=8)
    ax.tick_params(axis='y', labelsize=8)
    ax.yaxis.grid(True, color='white', linewidth=1.2, zorder=0)
    ax.set_axisbelow(True)
    for spine in ax.spines.values():
        spine.set_visible(False)
    ax.legend(fontsize=8)
    fig.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150)
    plt.close(fig)
    buf.seek(0)
    return buf

def make_ratio_chart(sessions, ratios, figsize=(7, 3)):
    fig, ax = plt.subplots(figsize=figsize)
    ax.set_facecolor('#F2F4F7')
    fig.patch.set_facecolor('#FFFFFF')
    colors = ['#1D6FDB', '#007A87', '#002B5C', '#E88C00']
    bars = ax.barh(sessions, ratios, color=colors[:len(sessions)], zorder=3)
    ax.set_title('Compression Ratio by Session (PO time vs. 3-person team equiv.)',
                 fontsize=10, fontweight='bold', color='#002B5C', pad=10)
    ax.set_xlabel('Compression Ratio (x)', fontsize=9)
    ax.xaxis.grid(True, color='white', linewidth=1.2, zorder=0)
    ax.set_axisbelow(True)
    for spine in ax.spines.values():
        spine.set_visible(False)
    for bar, val in zip(bars, ratios):
        ax.text(bar.get_width() + 1, bar.get_y() + bar.get_height()/2,
                f'{val}x', va='center', fontsize=9, color='#002B5C', fontweight='bold')
    ax.tick_params(axis='y', labelsize=8)
    ax.tick_params(axis='x', labelsize=8)
    fig.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150)
    plt.close(fig)
    buf.seek(0)
    return buf


# ── Cover page ─────────────────────────────────────────────────────────────────

def build_cover(doc):
    # Top colour block
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    para_shade(p, DARK_NAVY)
    run = p.add_run('\n\n')
    run.font.size = Pt(6)

    # Title band
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(0)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_shade(p2, DARK_NAVY)
    r1 = p2.add_run('Engineering Efficiency with AI Pair Programming')
    style_run(r1, size=22, bold=True, color=WHITE)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(4)
    p3.paragraph_format.space_after = Pt(0)
    para_shade(p3, DARK_NAVY)
    r2 = p3.add_run('A Case Study of ARIA — One Product Owner, One Claude')
    style_run(r2, size=15, bold=False, color=(0x90, 0xBC, 0xF8))

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.paragraph_format.space_before = Pt(6)
    p4.paragraph_format.space_after = Pt(0)
    para_shade(p4, DARK_NAVY)
    r3 = p4.add_run('Fintech SaaS · Two Products · Six Days · One Person')
    style_run(r3, size=11, italic=True, color=(0xCC, 0xDD, 0xFF))

    # Spacer
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(0)
    sp.paragraph_format.space_after = Pt(0)
    para_shade(sp, DARK_NAVY)
    sp.add_run('\n').font.size = Pt(6)

    add_hr(doc, color=BRAND_BLUE, thickness=20)

    # KPI row
    kpi_tbl = doc.add_table(rows=1, cols=4)
    kpi_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    kpis = [
        ('73', 'Commits', '6 days'),
        ('~12h', 'PO Active Time', 'total across 9 sessions'),
        ('2', 'Products Shipped', 'Advisor + Personal'),
        ('60–90×', 'Peak Compression', 'Session 1 ratio'),
    ]
    for i, (val, label, sub) in enumerate(kpis):
        cell = kpi_tbl.rows[0].cells[i]
        set_cell_bg(cell, LIGHT_GREY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(val + '\n')
        style_run(r, size=20, bold=True, color=BRAND_BLUE)
        r2 = p.add_run(label + '\n')
        style_run(r2, size=9, bold=True, color=DARK_NAVY)
        r3 = p.add_run(sub)
        style_run(r3, size=8, italic=True, color=(0x66, 0x77, 0x88))
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(8)
    for row in kpi_tbl.rows:
        for cell in row.cells:
            cell.width = Cm(4.0)

    add_hr(doc, color=BRAND_BLUE, thickness=8)

    # Meta info
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(10)
    meta.paragraph_format.space_after = Pt(4)
    r = meta.add_run('Author: sunder-vasudevan    ·    March 2026    ·    Version 1.0')
    style_run(r, size=9, italic=True, color=(0x55, 0x66, 0x77))

    audience = doc.add_paragraph()
    audience.alignment = WD_ALIGN_PARAGRAPH.CENTER
    audience.paragraph_format.space_before = Pt(2)
    audience.paragraph_format.space_after = Pt(16)
    r2 = audience.add_run('Audience: Technical peers · Engineering leads · Technical founders · Fintech product teams')
    style_run(r2, size=9, italic=True, color=(0x55, 0x66, 0x77))


# ── TOC page ───────────────────────────────────────────────────────────────────

def build_toc_page(doc):
    doc.add_page_break()
    add_heading(doc, 'Table of Contents', level=1, color=DARK_NAVY, size=16, before=0, after=10)
    add_hr(doc, color=BRAND_BLUE, thickness=12)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run()
    style_run(run, size=9, italic=True, color=(0x66, 0x77, 0x88))
    run.text = '(Update this field: select all → right-click → Update Field, or press Ctrl+A then F9)'
    insert_toc(doc)


# ── Section builders ───────────────────────────────────────────────────────────

def build_abstract(doc):
    doc.add_page_break()
    add_heading(doc, 'Abstract', level=1, color=DARK_NAVY, size=16)
    add_hr(doc, color=BRAND_BLUE)
    add_callout(doc,
        'This paper documents a real-world experiment: a single Product Owner building ARIA — '
        'a fintech SaaS platform comprising two distinct products (A-RiA Advisor Workbench and '
        'ARIA Personal) — in collaboration with Claude (Anthropic\'s AI assistant) as a '
        'persistent pair-programmer.',
        bg=LIGHT_GREY, color=DARK_NAVY, size=11, bold=False)
    add_body(doc,
        'Over 6 active coding days (2026-03-13 to 2026-03-18), the pair shipped both products from '
        'zero to market-ready state: a full-stack wealth management advisor workbench and a '
        'self-directed consumer finance app, sharing a common FastAPI backend. 9 sessions produced '
        '73 commits, spanning AI-powered features, a Monte Carlo simulation engine, JWT auth, cloud '
        'deployment (Render + Supabase + Vercel), 26 UI/UX polish fixes, production bug resolution, '
        'brand identity, and investor documentation — all from a single Product Owner with no '
        'engineering team.')
    add_body(doc,
        'The ARIA project demonstrates a sustained 20–40× compression ratio over equivalent '
        'traditional engineering effort, with individual sessions reaching 60–90× on active '
        'interaction time.')


def build_introduction(doc):
    doc.add_page_break()
    add_heading(doc, '1. Introduction — The Experiment', level=1, color=DARK_NAVY, size=16)
    add_hr(doc, color=BRAND_BLUE)
    add_body(doc,
        'ARIA was not a side project. It was built as a market-ready fintech product, targeting '
        'two distinct audiences: professional wealth management advisors (A-RiA) and self-directed '
        'retail investors (ARIA Personal). The architecture spans a shared FastAPI backend, two '
        'independent React/Vite frontends, PostgreSQL via Supabase, and live AI features powered '
        'by the Claude API — all deployed, all production.')
    add_heading(doc, 'The Central Question', level=2, color=MID_BLUE, size=12)
    add_callout(doc,
        'Can one Product Owner, working with Claude as a persistent full-session collaborator, '
        'build a fintech product to market-ready state in the time it would take a traditional '
        'team to write a sprint plan?',
        bg=(0xE8, 0xF0, 0xFE), color=DARK_NAVY, size=11, bold=True)
    add_heading(doc, 'What "Efficiency" Means Here', level=2, color=MID_BLUE, size=12)
    for bullet in [
        'Feature throughput per calendar day',
        'Architectural depth delivered without accruing uncontrolled debt',
        'Code quality signals visible in git history',
        'Documentation discipline sustained throughout',
        'Active Product Owner interaction time per feature shipped',
    ]:
        add_bullet(doc, bullet)
    add_heading(doc, 'Why Fintech Specifically Matters', level=2, color=MID_BLUE, size=12)
    add_body(doc,
        'Fintech products carry higher technical complexity than typical web apps: Monte Carlo '
        'simulation engines, probability scoring, inflation-adjusted projections, JWT auth, '
        'multi-tenant data isolation, and regulatory-adjacent data handling. This is not a todo '
        'app. The domain complexity makes the efficiency story more meaningful.')


def build_scope(doc):
    doc.add_page_break()
    add_heading(doc, '2. Project Scope — What Was Built', level=1, color=DARK_NAVY, size=16)
    add_hr(doc, color=BRAND_BLUE)

    add_heading(doc, 'A-RiA Advisor Workbench', level=2, color=MID_BLUE, size=13)
    add_body(doc, 'Professional tool for wealth management relationship managers:')
    for f in [
        'Client roster with segment, risk profile, portfolio, goals, and life events',
        'AI Copilot with client-specific context injection (Claude Sonnet 4.6)',
        'Morning Briefing: AI-generated per-client urgency analysis',
        'Meeting Prep Card: AI-generated agenda, risks, talking points, open questions',
        'Monte Carlo goal probability engine (1,000 simulation paths, inflation-adjusted)',
        'What-if scenario modeler with live sliders — Mode 1: monthly SIP, Mode 2: reverse-calculate required SIP',
        'Client Interaction Capture with urgency escalation',
        'Client Portal (client-facing view)',
        '26-item UI/UX polish pass (accessibility, mobile, skeleton loaders, animations, print styles)',
    ]:
        add_bullet(doc, f)

    add_heading(doc, 'ARIA Personal', level=2, color=MID_BLUE, size=13)
    add_body(doc, 'Consumer self-directed finance app:')
    for f in [
        'JWT registration and login (real user accounts, no hardcoding)',
        'Portfolio management with donut chart and holdings table',
        'Goal tracking with Monte Carlo probability rings',
        'Life events log',
        'AI Copilot in first-person voice ("your portfolio", not "the client")',
        'Shared backend infrastructure with advisor app',
    ]:
        add_bullet(doc, f)

    add_heading(doc, 'Technology Stack', level=2, color=MID_BLUE, size=12)
    add_table(doc,
        ['Layer', 'Technology'],
        [
            ['Frontend (Advisor)', 'React 18 + Vite + Tailwind CSS (navy palette)'],
            ['Frontend (Personal)', 'React 18 + Vite + Tailwind CSS (separate repo + Vercel)'],
            ['Backend', 'FastAPI (Python 3.11), SQLAlchemy, shared by both products'],
            ['Database', 'Supabase PostgreSQL (pooler, port 6543)'],
            ['Auth — Advisor', 'Frontend localStorage, hardcoded advisor credentials'],
            ['Auth — Personal', 'JWT (python-jose + bcrypt), real user accounts'],
            ['AI Features', 'Claude Sonnet 4.6 (Anthropic API)'],
            ['Simulation Engine', 'Custom Monte Carlo — 1,000 paths, log-normal returns, inflation-adjusted'],
            ['Deployment', 'Render (backend) + Vercel (frontend ×2) + Supabase (DB)'],
        ],
        col_widths=[5.5, 10.5]
    )

    add_heading(doc, 'Deployed URLs', level=2, color=MID_BLUE, size=12)
    add_table(doc,
        ['Product', 'URL', 'Status'],
        [
            ['A-RiA Advisor', 'https://a-ria.vercel.app', '✅ Live'],
            ['Backend API', 'https://aria-advisor.onrender.com', '✅ Live'],
            ['ARIA Personal', 'Vercel deploy pending', '🔄 Frontend built'],
        ],
        col_widths=[5, 8, 3]
    )


def build_methodology(doc):
    doc.add_page_break()
    add_heading(doc, '3. Methodology — How Efficiency Is Measured', level=1, color=DARK_NAVY, size=16)
    add_hr(doc, color=BRAND_BLUE)
    add_heading(doc, 'Primary Metric', level=2, color=MID_BLUE, size=12)
    add_body(doc,
        'Feature throughput per active day: how many meaningful product features were shipped '
        'per active coding day, compared to what a small team would typically ship in the same '
        'calendar window.')
    add_heading(doc, 'Comparison Baseline: Traditional 3-Person Startup Team', level=2, color=MID_BLUE, size=12)
    for b in [
        '1 frontend engineer + 1 backend engineer + 1 tech lead / PM',
        'Typical output in a 2-week sprint: 2–4 significant features to production',
        '15–25% of sprint consumed by ceremonies, reviews, coordination',
        'Architecture work treated as dedicated sprint scope, not background activity',
        'Deployment pipeline setup: 2–5 days',
        'Documentation: typically deferred',
    ]:
        add_bullet(doc, b)
    add_heading(doc, 'Data Sources', level=2, color=MID_BLUE, size=12)
    add_table(doc,
        ['Source', 'What It Measures'],
        [
            ['git log --shortstat', '73 commits across 6 active days, LOC per commit'],
            ['SESSION_LOG.md', 'All 9 sessions — goals, features shipped, estimated PO time'],
            ['INTERACTION_LOG.md', 'Formal session log from Session 4 onwards'],
            ['DECISION_LOG.md', 'Architecture decisions with explicit rationale'],
            ['PO self-assessment', '"2–3× faster" at project level — consistent with BzHub findings'],
        ],
        col_widths=[5, 11]
    )
    add_callout(doc,
        'Important: "Interaction time" measures only the Product Owner\'s active prompting and '
        'review time — not wall-clock time or Claude\'s generation time. It represents the '
        'human direction cost — the relevant denominator when evaluating how much PO judgment '
        'is needed to produce a given output.',
        bg=(0xFF, 0xF3, 0xCD), color=(0x7B, 0x4A, 0x00), size=10)


def build_timeline(doc):
    doc.add_page_break()
    add_heading(doc, '4. Timeline & Velocity Analysis', level=1, color=DARK_NAVY, size=16)
    add_hr(doc, color=BRAND_BLUE)

    add_heading(doc, 'Commit Activity by Date', level=2, color=MID_BLUE, size=12)
    add_table(doc,
        ['Date', 'Commits', 'Notes'],
        [
            ['2026-03-13', '8', 'Project init — full Phase 1 + Phase 2, PRD v1.0, AI features, 20 clients seeded'],
            ['2026-03-14', '3', 'ARIA rebrand, Monte Carlo engine (FEAT-501), session docs'],
            ['2026-03-15', '2', 'Session wrap docs + stability checkpoint'],
            ['2026-03-16', '7', 'Render + Supabase + Vercel deploy, Meeting Prep Card, Client Portal, HELP.md, v1.2'],
            ['2026-03-17', '15', 'FEAT-101 (add/edit client), FEAT-102/108/109 (onboarding wizard), mobile layout, AI error handling'],
            ['2026-03-18', '38', 'FEAT-404, 26 UI/UX fixes, ARIA Personal backend, What-if v2, goals/holdings CRUD, branding, whitepaper'],
        ],
        col_widths=[3, 2.5, 10.5]
    )

    # Commits chart
    chart_buf = make_bar_chart(
        ['Mar 13', 'Mar 14', 'Mar 15', 'Mar 16', 'Mar 17', 'Mar 18'],
        [8, 3, 2, 7, 15, 38],
        'Commits per Active Day — ARIA',
        ylabel='Commits'
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(chart_buf, width=Inches(6.0))

    add_callout(doc,
        '38 of 73 commits (52%) occurred on a single day (2026-03-18). The two peak days '
        '(Mar 17–18) account for 53 commits (73%). This intensity pattern — achievable without '
        'cognitive fatigue because Claude holds cross-session context — is the clearest signal '
        'of AI-amplified throughput.',
        bg=LIGHT_GREY, color=DARK_NAVY, size=10)

    add_heading(doc, 'Session-by-Session Breakdown', level=2, color=MID_BLUE, size=12)
    add_table(doc,
        ['Session', 'Date', 'PO Time', 'Key Output'],
        [
            ['1', '2026-03-13', '~90 min', 'Full Phase 1 + Phase 2 from zero: FastAPI, React, all AI features, 20 clients, PRD v1.0'],
            ['2', '2026-03-14', '~15 min', 'ARIA rebrand, Monte Carlo engine, tagline set'],
            ['3', '2026-03-15', '~20 min', 'Session wrap + phase sync'],
            ['4', '2026-03-16', '~3 hrs', 'Full cloud deploy (Render + Supabase + Vercel), Meeting Prep Card, Auth, Client Portal, HELP.md'],
            ['5', '2026-03-17', '~1h 45m', 'AI error handling, Morning Briefing redesign, URL migration'],
            ['6', '2026-03-17', '~30 min', 'Mobile-responsive layout (FEAT-407) — full mobile across all pages'],
            ['7', '2026-03-17', '~45 min', 'FEAT-101: Add/Edit Client — 7 new fields, PUT /clients/{id}, ClientForm.jsx'],
            ['8', '2026-03-17', '~30 min', 'FEAT-102/108/109: 4-tab onboarding wizard — risk questionnaire, portfolio, goals'],
            ['9', '2026-03-18', '~4 hrs', 'FEAT-404, 26 UX fixes, ARIA Personal backend, What-if v2, goals CRUD, login redesign, branding, whitepaper'],
            ['Total', '', '~12 hrs', '73 commits · 2 products · 6 active days'],
        ],
        col_widths=[1.5, 2.8, 2, 9.7]
    )

    add_heading(doc, 'The Day-One Sprint — Session 1 (2026-03-13)', level=2, color=MID_BLUE, size=12)
    add_callout(doc,
        'In approximately 90 minutes of active interaction, the pair shipped: FastAPI backend '
        'with all routers + 20 seeded clients, Claude API integration (Copilot, Briefing, '
        'Situation Summary), React + Vite frontend with all components, and PRD v1.0. A '
        'traditional 3-person team would plan this scope for a 4–6 week sprint.',
        bg=(0xE8, 0xF0, 0xFE), color=DARK_NAVY, size=10, bold=False)

    add_heading(doc, 'The Peak Day — Session 9 (2026-03-18)', level=2, color=MID_BLUE, size=12)
    add_body(doc, '38 commits. What shipped in one day:')
    add_table(doc,
        ['Feature', 'Scope / LOC'],
        [
            ['FEAT-404: Client Interaction Capture', 'New DB table, backend router, InteractionsPanel.jsx — 510 LOC'],
            ['UI/UX Batch 1–3 (26 fixes)', 'Skeleton loaders, touch targets, animations, sidebar collapse, print styles, aria-labels, lazy tabs'],
            ['ARIA Personal backend', 'JWT auth, 5 FastAPI routers, personal_models.py, DB migration — 753 LOC'],
            ['FEAT-503: What-if v2', 'Inflation-adjusted Monte Carlo + reverse-SIP calculator — 310 LOC'],
            ['Goals + Life Events + Holdings CRUD', 'Add/edit/delete across 3 entities, holding drawer, chart expand, NAV data — 1,292 LOC'],
            ['Login redesign — both apps', 'Split layout, dark navy gradient, slate-50 right, brand taglines'],
            ['ARiALogo / ARIALogo components', 'Round dot on ı, brand blue #1D6FDB, deployed across all pages'],
            ['ARIA Whitepaper + Executive Deck', 'ARIA_WHITEPAPER.md (5,868 words), Whitepaper.docx (116KB), Executive Deck (12 slides, 143KB)'],
        ],
        col_widths=[6, 10]
    )
    add_callout(doc,
        'A traditional 3-person team would allocate 6–10 engineer-weeks for this feature set.',
        bg=(0xFF, 0xF0, 0xF0), color=ROSE, size=10, bold=True)


def build_loc_analysis(doc):
    doc.add_page_break()
    add_heading(doc, '5. Lines of Code Analysis — Key Feature Commits', level=1, color=DARK_NAVY, size=16)
    add_hr(doc, color=BRAND_BLUE)
    add_table(doc,
        ['Date', 'Feature', 'Files', 'Insertions', 'Deletions'],
        [
            ['2026-03-13', 'Full Phase 1 + Phase 2 build', '~40', '~3,500', '0'],
            ['2026-03-16', 'Meeting Prep Card + Auth + Portal + HELP', '~15', '~900', '~50'],
            ['2026-03-17', 'FEAT-101: Add/Edit Client module', '13', '721', '8'],
            ['2026-03-17', 'FEAT-102/108/109: Onboarding wizard', '4', '980', '198'],
            ['2026-03-17', 'Mobile-responsive layout', '5', '300', '111'],
            ['2026-03-18', 'FEAT-404: Interaction Capture', '9', '510', '7'],
            ['2026-03-18', 'UI/UX Batch 1 (8 fixes)', '9', '633', '36'],
            ['2026-03-18', 'UI/UX Batch 2 (8 fixes)', '11', '122', '54'],
            ['2026-03-18', 'UI/UX Batch 3 (10 fixes)', '10', '247', '174'],
            ['2026-03-18', 'Add/Edit Goals + Life Events + Holdings', '11', '1,292', '232'],
            ['2026-03-18', 'FEAT-503: What-if v2', '5', '310', '70'],
            ['2026-03-18', 'ARIA Personal backend', '10', '753', '5'],
            ['2026-03-18', 'Login redesign + branding + docs batch', '11', '2,397', '26'],
        ],
        col_widths=[2.5, 7, 1.5, 2.5, 2.5]
    )
    add_callout(doc,
        'Estimated total net application LOC across both products: ~12,000–14,000 lines\n'
        '(Excludes lock files, migration scripts, documentation, and dependency artifacts)',
        bg=LIGHT_GREY, color=DARK_NAVY, size=10, bold=False)

    # LOC chart for key commits
    features = ['Phase 1+2\nBuild', 'FEAT-101\nClient', 'FEAT-102/108\nOnboarding', 'Mobile\nLayout',
                'FEAT-404\nInteractions', 'Goals+Holdings\nCRUD', 'FEAT-503\nWhat-if', 'Personal\nBackend']
    locs = [3500, 721, 980, 300, 510, 1292, 310, 753]
    chart_buf = make_bar_chart(features, locs, 'Net Insertions per Feature Commit', ylabel='Lines Inserted', figsize=(7.5, 3.5))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(chart_buf, width=Inches(6.5))


def build_comparison(doc):
    doc.add_page_break()
    add_heading(doc, '6. Comparative Analysis', level=1, color=DARK_NAVY, size=16)
    add_hr(doc, color=BRAND_BLUE)

    add_heading(doc, 'Full Project Window', level=2, color=MID_BLUE, size=12)
    add_table(doc,
        ['Dimension', '1 PO + Claude (6 days)', 'Est. 3-Person Team (6 days)'],
        [
            ['Products shipped', '2 (Advisor + Personal — shared backend)', '0–1 (planning + scaffolding phase)'],
            ['Production deployments', '3 (Render + Vercel ×2)', 'Setup in progress'],
            ['AI features live', '4 (Copilot, Briefing, Meeting Prep, Situation Summary)', '0 — AI integration is a dedicated sprint'],
            ['Custom simulation engine', '✅ Monte Carlo, 1,000 paths, inflation-adjusted', '0 — requires quant specialist time'],
            ['Auth systems', '2 (advisor frontend auth + Personal JWT)', '0–1'],
            ['UI/UX polish fixes', '26 systematic fixes', '4–6 (typical sprint capacity)'],
            ['Documentation maintained', 'Per-feature (enforced rule)', 'Deferred'],
            ['Investor materials', 'Whitepaper + DOCX + 12-slide PPTX', 'Not started'],
            ['Total PO active time', '~12 hours', 'N/A (team always staffed)'],
        ],
        col_widths=[5, 6, 5]
    )

    add_heading(doc, 'Compression Ratio by Session Type', level=2, color=MID_BLUE, size=12)
    add_table(doc,
        ['Session Type', 'PO Time', 'Est. 3-Person Equiv.', 'Ratio'],
        [
            ['Full product build from zero (Session 1)', '~90 min', '4–6 weeks', '60–90×'],
            ['Cloud deploy + live AI features (Session 4)', '~3 hrs', '2–3 weeks', '15–25×'],
            ['Feature sprint — 26 UX + 3 features + branding (Session 9)', '~4 hrs', '8–12 weeks', '25–40×'],
            ['Mobile layout pass (Session 6)', '~30 min', '3–5 days', '20–30×'],
        ],
        col_widths=[7.5, 2, 3.5, 3]
    )

    # Compression chart
    sessions = ['Session 1\n(Full build)', 'Session 4\n(Deploy + AI)', 'Session 9\n(Feature sprint)', 'Session 6\n(Mobile)']
    ratios = [75, 20, 32, 25]
    chart_buf = make_ratio_chart(sessions, ratios)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(chart_buf, width=Inches(6.0))

    add_heading(doc, 'Where the Gains Come From', level=2, color=MID_BLUE, size=12)
    gains = [
        ('Zero handoff overhead', 'Frontend, backend, DB schema, and AI prompt engineering all change in a single session without coordination latency.'),
        ('Persistent cross-session context', 'NOTES.md, PRD.md, SESSION_LOG.md, DECISION_LOG.md, and .claude/memory/ preserve full project state. Claude resumes with context a new hire would need 2–3 weeks to develop.'),
        ('On-demand full-stack + domain expertise', 'No single engineer is equally fluent in Monte Carlo simulation, FastAPI async patterns, React state lifting, Safari WebKit quirks, and Tailwind animation design. The pair covers all layers.'),
        ('Documentation as a zero-cost side-effect', 'HELP.md, PRD, RELEASE_NOTES, DECISION_LOG, whitepaper, and executive deck were all maintained throughout — not deferred.'),
        ('Parallel execution', 'Background agent pattern: documentation (whitepaper + PPTX) generated by a background agent while main thread continues feature work.'),
    ]
    for title, body in gains:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(title + ': ')
        style_run(r1, size=10.5, bold=True, color=DARK_NAVY)
        r2 = p.add_run(body)
        style_run(r2, size=10.5, color=DARK_TEXT)


def build_architecture(doc):
    doc.add_page_break()
    add_heading(doc, '7. Architecture Observations', level=1, color=DARK_NAVY, size=16)
    add_hr(doc, color=BRAND_BLUE)

    add_heading(doc, 'Two Products Without Architectural Debt', level=2, color=MID_BLUE, size=12)
    add_body(doc,
        'Building two products on a shared backend could easily produce a "frankenapp" — feature '
        'flags, tangled auth, mixed concerns. The ARIA architecture avoided this deliberately:')
    for b in [
        'Clean separation: separate repos, separate frontends, separate Vercel deployments',
        'Shared only what belongs shared: simulation engine, DB infrastructure, base models',
        'New personal routes under /personal/ prefix — zero pollution of existing advisor routes',
        'Nullable FKs for data model extension: personal_user_id added as nullable — zero risk to advisor data',
        'Each architectural decision documented in DECISION_LOG.md with explicit rationale',
    ]:
        add_bullet(doc, b)

    add_heading(doc, 'Known Debt — Openly Tracked', level=2, color=MID_BLUE, size=12)
    add_callout(doc,
        'ARIA Personal not yet deployed to Vercel (frontend built, awaiting deploy)\n'
        'Advisor auth is hardcoded (localStorage) — JWT upgrade planned for Phase 3\n'
        'DESIGN.md not yet written (design system verbally agreed, not codified)\n'
        '/help page not yet in-app (standing rule — flagged for next session)',
        bg=(0xFF, 0xF3, 0xCD), color=(0x7B, 0x4A, 0x00), size=10)

    add_heading(doc, 'The Simulation Engine', level=2, color=MID_BLUE, size=12)
    add_body(doc,
        'The Monte Carlo engine (backend/app/simulation.py) runs 1,000 simulation paths with '
        'log-normal return distribution per risk category, inflation adjustment (default 6% for '
        'Indian market context), binary search reverse-SIP calculation (Mode 2 what-if), and '
        'goal probability as P(final_value ≥ inflation_adjusted_target).')
    add_body(doc,
        'This is quantitative finance work — the kind that requires a dedicated quant engineer '
        'in a traditional team. It was built and iterated in two sessions.')

    add_heading(doc, 'Collaboration Patterns', level=2, color=MID_BLUE, size=12)
    add_table(doc,
        ['Task Category', 'Examples from ARIA'],
        [
            ['Full-stack feature design', 'DB model → FastAPI router → React component in one session'],
            ['Algorithm implementation', 'Monte Carlo engine, inflation adjustment, binary search SIP'],
            ['Architecture decisions', 'Shared backend vs. separate, JWT auth design, nullable FK migration'],
            ['Bug diagnosis', 'FastAPI 204 + implicit None = ASGI 500; Safari WebKit date onChange; React state remount'],
            ['Cross-browser compatibility', 'Safari date input → month/year selects (permanent codebase pattern)'],
            ['Design system', 'Brand blue #1D6FDB, probability pill colors, split-panel login layout'],
            ['Documentation', 'PRD, whitepaper, executive deck, HELP.md, DECISION_LOG'],
            ['Parallel workstreams', 'Background agent: whitepaper + PPTX while main thread ships features'],
        ],
        col_widths=[5, 11]
    )


def build_limitations(doc):
    doc.add_page_break()
    add_heading(doc, '8. Limitations & Honest Caveats', level=1, color=DARK_NAVY, size=16)
    add_hr(doc, color=BRAND_BLUE)
    caveats = [
        ('Partial time tracking', 'Interaction time was formally estimated for all sessions but not logged with stopwatch precision. Sessions 1–3 are approximate.'),
        ('"Interaction time" ≠ "engineering hours"', 'The metric measures active prompting and review time, not wall-clock time or Claude\'s generation time. It represents the human direction cost.'),
        ('No second human reviewer', 'There was no code review from a second engineer. Human review catches bugs and knowledge-sharing opportunities this model does not replicate.'),
        ('Greenfield advantage', 'Building from scratch allows architectural freedom that legacy codebases do not.'),
        ('Domain familiarity matters', 'The Product Owner had knowledge of Indian wealth management — INR figures, SIP mechanics, HNI segmentation. AI amplifies existing knowledge; it does not substitute for it.'),
        ('Multiplier depends on PO quality', 'A Product Owner who cannot read and evaluate code cannot safely use AI-generated code at this velocity.'),
    ]
    for title, body in caveats:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(3)
        r1 = p.add_run(title + ': ')
        style_run(r1, size=10.5, bold=True, color=DARK_NAVY)
        r2 = p.add_run(body)
        style_run(r2, size=10.5, color=DARK_TEXT)


def build_conclusions(doc):
    doc.add_page_break()
    add_heading(doc, '9. Conclusions & Recommendations', level=1, color=DARK_NAVY, size=16)
    add_hr(doc, color=BRAND_BLUE)

    add_heading(doc, 'Key Finding', level=2, color=MID_BLUE, size=13)
    add_callout(doc,
        'One capable Product Owner working with Claude as a persistent, full-session pair-programmer '
        'can build and deploy a production-grade fintech SaaS platform — two products, shared backend, '
        'live AI features, Monte Carlo simulation engine, JWT auth, cloud deployment — in 6 active '
        'coding days. The same scope would require a 3-person startup team 4–6 months.',
        bg=(0xE8, 0xF0, 0xFE), color=DARK_NAVY, size=11, bold=False)
    add_callout(doc,
        'For Session 1 specifically: ~90 minutes of active PO time produced a full-stack fintech '
        'application from zero — a scope a traditional team would plan for a 4–6 week sprint. '
        'This represents a ~60–90× compression ratio on human direction time.',
        bg=LIGHT_GREY, color=DARK_NAVY, size=11, bold=True)

    add_heading(doc, 'Structural Advantages', level=2, color=MID_BLUE, size=12)
    for adv in [
        'Zero handoff latency across the full stack',
        'Persistent cross-session context with zero onboarding overhead',
        'On-demand expertise across frontend, backend, DB, AI, and quantitative finance',
        'Documentation as a zero-marginal-cost side-effect',
        'Parallel execution via background agents',
    ]:
        add_bullet(doc, adv)

    add_heading(doc, 'Recommended Workflow Patterns', level=2, color=MID_BLUE, size=12)
    recs = [
        ('Invest in context management', 'NOTES.md + PRD + SESSION_LOG + DECISION_LOG is the highest-leverage infrastructure investment. Eliminates the 15–30 minute catch-up cost at every session start.'),
        ('Enforce documentation as a session rule', 'HELP.md, release notes, and decision log updated per feature. This discipline is only sustainable because Claude handles the mechanical writing.'),
        ('Keep the engineer in the architecture seat', 'Product direction, data model decisions, and quality review remain the PO\'s responsibility. AI amplifies execution; it does not replace judgment.'),
        ('Use the parallel execution pattern', 'Background agents for documentation, infrastructure, or parallel feature tracks while the main thread continues core work.'),
        ('Log interaction time from day one', 'The clearest evidence of efficiency is measured PO time per feature shipped. Start logging from the first session.'),
    ]
    for i, (title, body) in enumerate(recs, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(3)
        r1 = p.add_run(f'{i}. {title}: ')
        style_run(r1, size=10.5, bold=True, color=DARK_NAVY)
        r2 = p.add_run(body)
        style_run(r2, size=10.5, color=DARK_TEXT)

    add_heading(doc, 'Where This Model Works Best', level=2, color=MID_BLUE, size=12)
    for w in [
        'Greenfield fintech feature development with clear domain requirements',
        'Full-stack features requiring frontend + backend + DB + AI in a single session',
        'Algorithm and simulation work where the domain logic is well-specified',
        'Systematic UI/UX polish passes — AI addresses design debt at speed',
        'Documentation and investor materials generation in parallel with feature work',
    ]:
        add_bullet(doc, w)

    add_heading(doc, 'Where Human Judgment Remains Essential', level=2, color=MID_BLUE, size=12)
    for w in [
        'Product and market positioning decisions',
        'Security-sensitive code (auth, RLS, data access policies)',
        'Code review and quality gating of AI output',
        'Domain expertise — knowing what "correct" looks like in fintech context',
        'Recognising when generated code is plausible but subtly wrong',
    ]:
        add_bullet(doc, w)


def build_appendices(doc):
    doc.add_page_break()
    add_heading(doc, 'Appendix A — Full Commit Log (Feature Commits)', level=1, color=DARK_NAVY, size=16)
    add_hr(doc, color=BRAND_BLUE)
    add_table(doc,
        ['Date', 'Session', 'Commit Summary', 'Files', 'Ins', 'Del'],
        [
            ['2026-03-13', '1', 'Full Phase 1 + Phase 2 — FastAPI, React, AI, 20 clients seeded', '~40', '~3,500', '0'],
            ['2026-03-14', '2', 'FEAT-501: Monte Carlo goal projection engine', '~3', '~200', '~10'],
            ['2026-03-14', '2', 'Rebrand: RIA Advisor Copilot → ARIA', '~5', '~30', '~20'],
            ['2026-03-16', '4', 'Meeting Prep Card, Advisor Login, Client Login, Client Portal', '~15', '~900', '~50'],
            ['2026-03-16', '4', 'HELP.md, v1.2 version in UI', '~3', '~350', '0'],
            ['2026-03-17', '5', 'Redesign Morning Briefing — structured layout + all-clear', '1', '80', '14'],
            ['2026-03-17', '6', 'Mobile-responsive layout — all pages (FEAT-407)', '5', '300', '111'],
            ['2026-03-17', '7', 'FEAT-101: Add/Edit Client — 7 fields, POST/PUT /clients', '13', '721', '8'],
            ['2026-03-17', '8', 'FEAT-102/108/109: 4-tab onboarding wizard', '4', '980', '198'],
            ['2026-03-18', '9', 'FEAT-404: Client Interaction Capture', '9', '510', '7'],
            ['2026-03-18', '9', 'UI/UX Batch 1 — 8 fixes (skeleton, touch, tabular nums)', '9', '633', '36'],
            ['2026-03-18', '9', 'UI/UX Batch 2 — 8 fixes (prefetch, shadows, text-wrap)', '11', '122', '54'],
            ['2026-03-18', '9', 'UI/UX Batch 3 — 10 fixes (animation, radius, sidebar, print)', '10', '247', '174'],
            ['2026-03-18', '9', 'Add/Edit/Delete Goals + Life Events + Holdings + NAV', '11', '1,292', '232'],
            ['2026-03-18', '9', 'FEAT-503: What-if Scenario v2 — Monte Carlo + reverse SIP', '5', '310', '70'],
            ['2026-03-18', '9', 'ARIA Personal: backend — JWT auth, 5 routers, personal models', '10', '753', '5'],
            ['2026-03-18', '9', 'Login redesign both apps + ARiALogo + brand identity', '~15', '~600', '~250'],
            ['2026-03-18', '9', 'ARIA Whitepaper + DOCX + Executive Deck (12 slides)', '3', '2,000+', '0'],
        ],
        col_widths=[2.5, 1.5, 7.5, 1.5, 1.5, 1.5]
    )

    doc.add_page_break()
    add_heading(doc, 'Appendix B — Feature Inventory', level=1, color=DARK_NAVY, size=16)
    add_hr(doc, color=BRAND_BLUE)
    add_heading(doc, 'A-RiA Advisor Workbench (v1.2)', level=2, color=MID_BLUE, size=12)
    add_table(doc,
        ['Feature', 'Status'],
        [
            ['Client List — grouped by urgency / list toggle', '✅ Live'],
            ['Client 360 — portfolio, goals, life events, interactions, AI tabs', '✅ Live'],
            ['AI Copilot (client-context-aware chat)', '✅ Live'],
            ['Morning Briefing (AI-generated per-advisor briefing)', '✅ Live'],
            ['Situation Summary (per-client AI analysis)', '✅ Live'],
            ['Meeting Prep Card (FEAT-308)', '✅ Live'],
            ['Monte Carlo goal projection engine', '✅ Live'],
            ['What-if Scenario v2 (inflation-adjusted, reverse SIP)', '✅ Live'],
            ['Add/Edit/Delete Goals inline', '✅ Live'],
            ['Add/Edit/Delete Life Events', '✅ Live'],
            ['Edit Holdings inline + detail drawer', '✅ Live'],
            ['Client Interaction Capture + urgency flags', '✅ Live'],
            ['Add/Edit Client (FEAT-101) — 7 fields', '✅ Live'],
            ['Risk Questionnaire (FEAT-102)', '✅ Live'],
            ['Client Onboarding Wizard (FEAT-108/109)', '✅ Live'],
            ['Mobile-responsive layout (FEAT-407)', '✅ Live'],
            ['Advisor Login + Client Login + Client Portal', '✅ Live'],
            ['HELP.md guide', '✅ Available (in-app page pending)'],
        ],
        col_widths=[12, 4]
    )
    add_heading(doc, 'ARIA Personal (v0.1)', level=2, color=MID_BLUE, size=12)
    add_table(doc,
        ['Feature', 'Status'],
        [
            ['JWT registration + login (real user accounts)', '✅ Built'],
            ['Portfolio management + donut chart', '✅ Built'],
            ['Goal tracking with probability rings', '✅ Built'],
            ['Life events log', '✅ Built'],
            ['AI Copilot (first-person voice)', '✅ Built'],
            ['Vercel deployment', '🔄 Pending'],
        ],
        col_widths=[12, 4]
    )

    doc.add_page_break()
    add_heading(doc, 'Appendix C — Session Log', level=1, color=DARK_NAVY, size=16)
    add_hr(doc, color=BRAND_BLUE)
    add_table(doc,
        ['#', 'Date', 'Goal', 'PO Time', 'Key Commits'],
        [
            ['1', '2026-03-13', 'Full Phase 1 + Phase 2 build', '~90 min', '439245c'],
            ['2', '2026-03-14', 'ARIA rebrand + Monte Carlo engine', '~15 min', '3442209, 9fbb726'],
            ['3', '2026-03-15', 'Session wrap + phase sync', '~20 min', '5dfb996'],
            ['4', '2026-03-16', 'Cloud deploy + Meeting Prep + auth', '~3 hrs', '4b359f7, eada25c'],
            ['5', '2026-03-17', 'AI error handling + Briefing redesign + URL migration', '~1h 45m', '335c8e4 → 005307d'],
            ['6', '2026-03-17', 'Mobile layout (FEAT-407)', '~30 min', 'c7fad1b'],
            ['7', '2026-03-17', 'FEAT-101: Add/Edit Client', '~45 min', '01d1af3'],
            ['8', '2026-03-17', 'FEAT-102/108/109: Onboarding wizard', '~30 min', '87fb965'],
            ['9', '2026-03-18', 'Interactions, 26 UX, Personal, What-if v2, branding, whitepaper', '~4 hrs', 'f6250ca → 70e6951'],
            ['Total', '', '', '~12 hrs active', '73 commits'],
        ],
        col_widths=[1, 2.5, 6, 2, 4.5]
    )

    doc.add_page_break()
    add_heading(doc, 'Appendix D — Stack Reference', level=1, color=DARK_NAVY, size=16)
    add_hr(doc, color=BRAND_BLUE)
    add_table(doc,
        ['Component', 'Technology'],
        [
            ['Backend framework', 'FastAPI (Python 3.11)'],
            ['ORM', 'SQLAlchemy'],
            ['Frontend', 'React 18 + Vite + Tailwind CSS'],
            ['AI model', 'Claude Sonnet 4.6 (Anthropic API)'],
            ['Database', 'Supabase PostgreSQL (pooler port 6543)'],
            ['Personal auth', 'JWT (python-jose + bcrypt)'],
            ['Deployment: backend', 'Render (free tier, auto-deploy)'],
            ['Deployment: frontend', 'Vercel (auto-deploy on push to main)'],
            ['Simulation engine', 'Custom Python Monte Carlo (1,000 paths, log-normal returns)'],
        ],
        col_widths=[5.5, 10.5]
    )


# ── Footer & page numbers ──────────────────────────────────────────────────────

def add_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.clear()
    r1 = p.add_run('ARIA — Engineering Efficiency Whitepaper   ·   sunder-vasudevan   ·   March 2026   ·   Page ')
    style_run(r1, size=8, italic=True, color=(0x88, 0x99, 0xAA))
    # Page number field
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run_el = OxmlElement('w:r')
    run_el.append(fldChar1)
    run_el.append(instrText)
    run_el.append(fldChar2)
    p._p.append(run_el)


# ── Page setup ─────────────────────────────────────────────────────────────────

def setup_page(doc):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)


# ── Main ───────────────────────────────────────────────────────────────────────

def main():
    doc = Document()
    setup_page(doc)

    # Override default paragraph style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)

    build_cover(doc)
    build_toc_page(doc)
    build_abstract(doc)
    build_introduction(doc)
    build_scope(doc)
    build_methodology(doc)
    build_timeline(doc)
    build_loc_analysis(doc)
    build_comparison(doc)
    build_architecture(doc)
    build_limitations(doc)
    build_conclusions(doc)
    build_appendices(doc)
    add_footer(doc)

    out_dir = os.path.dirname(os.path.abspath(__file__))
    out_path = os.path.join(out_dir, 'ARIA_Efficiency_Whitepaper.docx')
    doc.save(out_path)
    print(f'✅  Saved: {out_path}')

if __name__ == '__main__':
    main()
